from django.conf import settings
from django.http import HttpResponse
from rest_framework import status, viewsets
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework.permissions import IsAuthenticated
from rest_framework.exceptions import NotFound
from docx import Document
import google.generativeai as genai
import io
import base64
from PIL import Image
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document
import logging
import traceback
import json
import os
from docx.shared import Inches
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image as RLImage
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.units import inch


logger = logging.getLogger(__name__)

from .models import Report
from .serializers import ReportSerializer, ReportGenerateSerializer


class ReportViewSet(viewsets.ReadOnlyModelViewSet):
    """ViewSet for viewing reports"""
    serializer_class = ReportSerializer
    permission_classes = [IsAuthenticated]
    
    def get_queryset(self):
        return Report.objects.filter(user=self.request.user)


class GenerateReportView(APIView):
    """Generate a medical report from an image"""
    permission_classes = [IsAuthenticated]
    
    def post(self, request):
        try:
            serializer = ReportGenerateSerializer(data=request.data)
            if not serializer.is_valid():
                logger.error("Serializer validation error: %s", serializer.errors)
                return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

            image_file = serializer.validated_data['image']
            language_code = serializer.validated_data.get('language', 'en')
            
            # Debug logging
            logger.info("API Key present: %s", bool(settings.GOOGLE_AI_API_KEY))
            logger.info("API Key length: %d", len(settings.GOOGLE_AI_API_KEY) if settings.GOOGLE_AI_API_KEY else 0)
            logger.info("Image file type: %s", type(image_file))
            logger.info("Image file name: %s", getattr(image_file, 'name', 'No name'))
            
            try:
                # Configure the Gemini AI model
                genai.configure(api_key=settings.GOOGLE_AI_API_KEY)
                
                # Get the model - using the newer version
                model = genai.GenerativeModel('gemini-1.5-flash')
                
                # Process the image
                try:
                    image_pil = Image.open(image_file)
                    logger.info("Image opened successfully. Size: %s", image_pil.size)
                except Exception as e:
                    logger.error("Error opening image: %s", str(e))
                    return Response({
                        'error': f"Error processing image: {str(e)}"
                    }, status=status.HTTP_400_BAD_REQUEST)
                
                # Generate report using Gemini
                try:
                    logger.info("Sending request to Gemini API...")
                    prompt = f"""Analyze this medical image and generate a detailed report in {language_code}. 
                    Please structure your response exactly as follows:

                    DIAGNOSIS:
                    [Write the main diagnosis or findings here]

                    DETAILS:
                    [Write detailed analysis here]

                    ACCURACY:
                    [Write a number between 0 and 1 indicating confidence level]

                    RECOMMENDATIONS:
                    [Write each recommendation on a new line with a bullet point]
                    """
                    
                    response = model.generate_content([prompt, image_pil])
                    logger.info("Received response from Gemini API")
                    
                    # Parse the response into structured data
                    report_text = response.text
                    logger.info("Raw AI response: %s", report_text)
                    
                    # Initialize default values
                    diagnosis = "No diagnosis available"
                    details = "No detailed analysis available"
                    accuracy = 0.8  # Default value
                    recommendations = ["No specific recommendations available"]
                    
                    # Split the response into sections
                    sections = report_text.split('\n\n')
                    current_section = None
                    
                    for section in sections:
                        section = section.strip()
                        if not section:
                            continue
                            
                        if section.upper().startswith('DIAGNOSIS:'):
                            current_section = 'diagnosis'
                            diagnosis = section.split(':', 1)[1].strip() if ':' in section else section.strip()
                        elif section.upper().startswith('DETAILS:'):
                            current_section = 'details'
                            details = section.split(':', 1)[1].strip() if ':' in section else section.strip()
                        elif section.upper().startswith('ACCURACY:'):
                            current_section = 'accuracy'
                            try:
                                accuracy_text = section.split(':', 1)[1].strip() if ':' in section else section.strip()
                                accuracy = float(accuracy_text)
                                # Ensure accuracy is between 0 and 1
                                accuracy = max(0.0, min(1.0, accuracy))
                            except ValueError:
                                accuracy = 0.8
                        elif section.upper().startswith('RECOMMENDATIONS:'):
                            current_section = 'recommendations'
                            recs = section.split(':', 1)[1].strip() if ':' in section else section.strip()
                            recommendations = [r.strip().lstrip('•- ') for r in recs.split('\n') if r.strip()]
                            if not recommendations:
                                recommendations = ["No specific recommendations available"]
                    
                    logger.info("Parsed sections - Diagnosis: %s, Details: %s, Accuracy: %s, Recommendations: %s",
                              diagnosis, details, accuracy, recommendations)
                    
                except Exception as e:
                    logger.error("Error from Gemini API: %s", str(e))
                    logger.error("Full traceback: %s", traceback.format_exc())
                    return Response({
                        'error': f"Error generating report with AI: {str(e)}"
                    }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
                
                # Create a new report in the database
                try:
                    report = Report.objects.create(
                        user=request.user,
                        image=image_file,
                        diagnosis=diagnosis,
                        details=details,
                        accuracy=accuracy,
                        recommendations=json.dumps(recommendations),
                        language=language_code
                    )
                    logger.info("Report created in database with ID: %s", report.id)
                except Exception as e:
                    logger.error("Error saving report to database: %s", str(e))
                    return Response({
                        'error': f"Error saving report: {str(e)}"
                    }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
                
               # Generate PDF report
                try:
                    buffer = io.BytesIO()
                    doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=50, leftMargin=50, topMargin=60, bottomMargin=40)
                    elements = []
                    styles = getSampleStyleSheet()
                    styles.add(ParagraphStyle(name='CenterTitle', parent=styles['Heading1'], alignment=TA_CENTER, fontSize=20, spaceAfter=20))
                    styles.add(ParagraphStyle(name='SectionHeader', parent=styles['Heading2'], fontSize=14, spaceAfter=8, spaceBefore=16))
                    styles.add(ParagraphStyle(name='NormalText', parent=styles['Normal'], fontSize=12, leading=16))

                    # Logo
                    current_dir = os.path.dirname(os.path.abspath(__file__))
                    logo_path = os.path.join(current_dir, '..', 'med', 'logo.png')
                    logo_path = os.path.abspath(logo_path)
                    logger.info(f"Tentative d'insertion du logo. Chemin utilisé : {logo_path}")
                    if os.path.exists(logo_path):
                        try:
                            img = RLImage(logo_path, width=2.0*inch, height=1.0*inch)
                            img.hAlign = 'CENTER'
                            elements.append(img)
                        except Exception as img_exc:
                            logger.error(f"Erreur lors de l'insertion du logo dans le PDF : {img_exc}")
                    else:
                        logger.warning(f"Logo not found at {logo_path}, skipping logo in PDF.")

                    # Title
                    elements.append(Paragraph("Medical Report", styles['CenterTitle']))

                    # Diagnosis
                    elements.append(Paragraph("Diagnosis:", styles['SectionHeader']))
                    elements.append(Paragraph(diagnosis, styles['NormalText']))

                    # Details
                    elements.append(Paragraph("Details:", styles['SectionHeader']))
                    elements.append(Paragraph(details.replace('\n', '<br/>'), styles['NormalText']))

                    # Accuracy
                    elements.append(Paragraph("Accuracy:", styles['SectionHeader']))
                    elements.append(Paragraph(f"{accuracy * 100:.1f}%", styles['NormalText']))

                    # Recommendations
                    elements.append(Paragraph("Recommendations:", styles['SectionHeader']))
                    for rec in recommendations:
                        elements.append(Paragraph(f"• {rec}", styles['NormalText']))

                    doc.build(elements)
                    pdf_buffer = buffer.getvalue()
                    logger.info("PDF report generated successfully")

                except Exception as e:
                    logger.error(f"Error generating PDF report: {e}")
                    pdf_buffer = None  # Optionnel : pour éviter l'utilisation de pdf_buffer plus loin si l'erreur se produit


                # Generate Word document
                try:
                    doc = Document()

                    # Add logo in header
                    logo_path = os.path.join(settings.BASE_DIR, 'med', 'logo.png')
                    logger.info(f"Tentative d'insertion du logo Word. Chemin utilisé : {logo_path}")
                    doc.sections[0].header.is_linked_to_previous = False
                    header = doc.sections[0].header
                    header_paragraph = header.paragraphs[0]
                    run = header_paragraph.add_run()
                    if os.path.exists(logo_path):
                        run.add_picture(str(logo_path), width=Inches(1.5))
                    else:
                        logger.warning(f"Logo not found at {logo_path}, skipping logo in Word doc.")

                    doc.add_heading('Medical Report', 0)
                    doc.add_heading('Diagnosis', level=1)
                    doc.add_paragraph(diagnosis)
                    doc.add_heading('Details', level=1)
                    doc.add_paragraph(details)
                    doc.add_heading('Accuracy', level=1)
                    doc.add_paragraph(f"{accuracy * 100:.1f}%")
                    doc.add_heading('Recommendations', level=1)
                    for rec in recommendations:
                        doc.add_paragraph(f"• {rec}", style='List Bullet')

                    word_buffer = io.BytesIO()
                    doc.save(word_buffer)
                    word_buffer.seek(0)
                    logger.info("Word document generated successfully")

                except Exception as e:
                    logger.error(f"Error generating Word document: {e}")
                    word_buffer = None  # Pareil : pour prévenir les erreurs plus loin

                
                # Return the report data
                response_data = {
                    'report_id': report.id,
                    'diagnosis': diagnosis,
                    'details': details,
                    'accuracy': accuracy,
                    'recommendations': recommendations,
                }
                if pdf_buffer:
                    response_data['pdf'] = base64.b64encode(pdf_buffer).decode('utf-8')
                if word_buffer:
                    response_data['docx'] = base64.b64encode(word_buffer.getvalue()).decode('utf-8')
                return Response(response_data, status=status.HTTP_201_CREATED)
                
            except Exception as e:
                logger.error("Unexpected error: %s", str(e))
                logger.error("Full traceback: %s", traceback.format_exc())
                return Response({
                    'error': f"Unexpected error: {str(e)}"
                }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
                
        except Exception as e:
            logger.error("Top-level error: %s", str(e))
            logger.error("Full traceback: %s", traceback.format_exc())
            return Response({
                'error': f"Server error: {str(e)}"
            }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)